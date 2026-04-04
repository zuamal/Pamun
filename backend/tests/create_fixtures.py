"""One-time script to generate binary test fixtures (docx, pdf).
Run once: python tests/create_fixtures.py
"""
from pathlib import Path

FIXTURES = Path(__file__).parent / "fixtures"


def create_docx() -> None:
    import docx

    doc = docx.Document()
    doc.add_paragraph("Users must be able to log in with email and password.")
    doc.add_paragraph("The system shall lock accounts after 5 failed attempts.")
    doc.add_paragraph("After login, users are redirected to the dashboard.")
    doc.save(FIXTURES / "small_sample.docx")
    print("Created small_sample.docx")


def create_pdf() -> None:
    from fpdf import FPDF  # type: ignore[import-untyped]

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Users must be able to log in with email and password.", ln=True)
    pdf.cell(0, 10, "The system shall lock accounts after 5 failed attempts.", ln=True)
    pdf.cell(0, 10, "After login, users are redirected to the dashboard.", ln=True)
    pdf.output(str(FIXTURES / "small_sample.pdf"))
    print("Created small_sample.pdf")


if __name__ == "__main__":
    create_docx()
    create_pdf()
